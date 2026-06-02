from __future__ import annotations

from pathlib import Path
from typing import Any

from app.infrastructure.parsers.base import ParsedDocument, ParsedPage, ParsedTable


class DocxParser:
    file_type = "docx"

    def parse(self, path: str | Path) -> ParsedDocument:
        from docx import Document

        path = Path(path)
        document = Document(path)

        pages: list[ParsedPage] = []
        current_section = None
        text_blocks = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                current_section = text
            text_blocks.append(text)

        if text_blocks:
            pages.append(
                ParsedPage(
                    page_number=None,
                    text="\n".join(text_blocks),
                    section_name=current_section,
                )
            )

        tables: list[ParsedTable] = []
        for table_index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            headers = [
                value if value else f"column_{index + 1}"
                for index, value in enumerate(rows[0])
            ]
            parsed_rows: list[dict[str, Any]] = []
            for row in rows[1:]:
                parsed_rows.append(
                    {
                        headers[index] if index < len(headers) else f"column_{index + 1}": value
                        for index, value in enumerate(row)
                    }
                )
            tables.append(
                ParsedTable(
                    source_name=f"docx_table_{table_index}",
                    page_number=None,
                    sheet_name=None,
                    rows=parsed_rows,
                )
            )

        metadata = {"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)}
        return ParsedDocument(file_type=self.file_type, pages=pages, tables=tables, metadata=metadata)
