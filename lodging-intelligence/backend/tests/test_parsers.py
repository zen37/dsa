from __future__ import annotations

import csv

from app.infrastructure.parsers.csv_parser import CsvParser
from app.infrastructure.parsers.docx_parser import DocxParser
from app.infrastructure.parsers.pdf_parser import PdfParser
from app.infrastructure.parsers.xlsx_parser import XlsxParser


def test_pdf_parser_returns_page_level_text(tmp_path):
    import fitz

    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hotel parking includes 100 spaces.")
    doc.save(path)
    doc.close()

    parsed = PdfParser().parse(path)

    assert parsed.file_type == "pdf"
    assert parsed.pages[0].page_number == 1
    assert "parking" in parsed.pages[0].text


def test_docx_parser_returns_paragraphs_and_tables(tmp_path):
    from docx import Document

    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Property Overview", level=1)
    doc.add_paragraph("The hotel has 132 keys.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "ADR"
    table.cell(1, 1).text = "251.89"
    doc.save(path)

    parsed = DocxParser().parse(path)

    assert "132 keys" in parsed.pages[0].text
    assert parsed.tables[0].rows[0]["Metric"] == "ADR"


def test_xlsx_parser_returns_sheet_level_table_data(tmp_path):
    from openpyxl import Workbook

    path = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Operating Statement"
    ws.append(["Line", "Amount"])
    ws.append(["Rooms Revenue", 9708984])
    wb.save(path)

    parsed = XlsxParser().parse(path)

    assert parsed.tables[0].sheet_name == "Operating Statement"
    assert parsed.tables[0].rows[0]["Line"] == "Rooms Revenue"
    assert parsed.pages[0].sheet_name == "Operating Statement"


def test_csv_parser_returns_table_data(tmp_path):
    path = tmp_path / "sample.csv"
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["hotel", "keys"])
        writer.writeheader()
        writer.writerow({"hotel": "Example Inn", "keys": "120"})

    parsed = CsvParser().parse(path)

    assert parsed.tables[0].rows[0]["hotel"] == "Example Inn"
    assert "Example Inn" in parsed.pages[0].text
